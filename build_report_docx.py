#!/usr/bin/env python3
"""
Genera Devotio_Reporte_Validacion.docx desde el markdown de hallazgos.
Uso: python3 build_report_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta de colores ──────────────────────────────────────────────────────────
C_BRAND_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # azul muy oscuro — encabezados
C_BRAND_MID    = RGBColor(0x16, 0x21, 0x3E)   # azul medio — subtítulos
C_ACCENT       = RGBColor(0x0F, 0x3A, 0x75)   # azul acento — bordes tabla
C_RED          = RGBColor(0xC0, 0x39, 0x2B)   # rojo crítico
C_ORANGE       = RGBColor(0xD3, 0x68, 0x00)   # naranja alto
C_GREEN        = RGBColor(0x1A, 0x7A, 0x40)   # verde ok
C_YELLOW_BG    = RGBColor(0xFF, 0xF9, 0xE6)   # fondo alerta suave
C_TABLE_HEADER = RGBColor(0x0F, 0x3A, 0x75)   # fondo encabezado tabla
C_TABLE_ROW    = RGBColor(0xF0, 0xF4, 0xFA)   # fila alternada tabla
C_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT         = RGBColor(0x1C, 0x1C, 0x1C)
C_MUTED        = RGBColor(0x55, 0x55, 0x55)


def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = rgb_hex(color)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_side="bottom", size=6, color="0F3A75"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{border_side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)


def add_run(para, text, bold=False, italic=False, color=None, size=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    # thick left border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "thick")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), rgb_hex(C_ACCENT))
    pBdr.append(left)
    pPr.append(pBdr)
    add_run(p, text, bold=True, color=C_BRAND_DARK, size=14)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, color=C_BRAND_MID, size=12)
    return p


def body(doc, text, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, color=color or C_TEXT, italic=italic, size=10)
    return p


def bullet(doc, text, level=0, color=None, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, color=color or C_TEXT, size=10)
        add_run(p, text, color=color or C_TEXT, size=10)
    else:
        add_run(p, text, color=color or C_TEXT, size=10)
    return p


def add_table(doc, headers, rows, col_widths=None, zebra=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_TABLE_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = C_TABLE_ROW if (zebra and r_idx % 2 == 0) else C_WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Color coding for specific keywords
            text = str(val)
            color = C_TEXT
            bold = False
            if text in ("🔴 CRÍTICO", "CRÍTICO"):
                color = C_RED
                bold = True
            elif text in ("🟠 ALTO", "ALTO"):
                color = C_ORANGE
                bold = True
            elif text in ("✅", "0 — Fix confirmado ✅", "89.8%", "77.3%"):
                color = C_GREEN
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = color
            run.bold = bold

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def add_status_badge(para, label, text_color=None):
    run = para.add_run(f" {label} ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = text_color or C_WHITE
    run.font.name = "Calibri"


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_callout(doc, text, icon="ℹ️"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "D36800")
    pBdr.append(left)
    pPr.append(pBdr)
    add_run(p, f"{icon}  {text}", italic=True, color=C_MUTED, size=9.5)
    return p


# ── Build document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover block ──────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(10)
    cover.paragraph_format.space_after = Pt(4)
    add_run(cover, "REPORTE DE VALIDACIÓN", bold=True, color=C_BRAND_DARK, size=20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    add_run(sub, "Sistema Stripe Automatizado — Devotio Rewards", bold=False, color=C_BRAND_MID, size=13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    add_run(meta, "Período: 29 de abril – 5 de mayo, 2026     |     Emitido: 6 de mayo, 2026     |     Blue Phoenix Lab",
            color=C_MUTED, size=9)

    add_divider(doc)

    # ── 1. Resumen Ejecutivo ─────────────────────────────────────────────────
    heading1(doc, "1. Resumen Ejecutivo")
    body(doc, (
        "Se realizó validación cruzada entre el reporte manual de Stripe Dashboard y el reporte automático "
        "generado por N8N. Se analizan dos ventanas: el período completo del 29 de abril al 5 de mayo y, "
        "con énfasis especial, el período del 1 al 5 de mayo para verificar posibles cobros duplicados."
    ))

    heading2(doc, "Resultados — Mayo 1–5 (ventana principal)")
    add_table(doc,
        ["Métrica", "Valor"],
        [
            ["Total transacciones en Stripe (manual)", "78"],
            ["Transacciones exitosas (Paid)", "34 — $3,886.30"],
            ["Transacciones capturadas correctamente", "69 de 78"],
            ["Cobertura del sistema", "89.8%"],
            ["Monto matcheado correctamente", "$3,491.55"],
            ["Pagos exitosos sin capturar", "5 — $394.75"],
            ["Bug de doble conteo (invoice.paid)", "0 — Fix confirmado ✅"],
        ],
        col_widths=[9, 6],
    )

    heading2(doc, "Resultados — Abril 29–Mayo 5 (período extendido)")
    add_table(doc,
        ["Métrica", "Valor"],
        [
            ["Total transacciones en Stripe", "124"],
            ["Transacciones exitosas (Paid)", "65 — $7,102.65"],
            ["Cobertura del sistema", "77.3%"],
            ["Monto matcheado correctamente", "$5,486.85"],
            ["Pagos exitosos sin capturar", "14 — $1,615.80"],
        ],
        col_widths=[9, 6],
    )

    add_callout(doc,
        "La cobertura baja de 89.8% a 77.3% al incluir el 29–30 de abril porque ese fin de semana "
        "tiene 9 pagos adicionales sin capturar ($1,221.05). Ver Sección 4.",
        icon="⚠️"
    )

    add_divider(doc)

    # ── 2. Estado del Sistema ────────────────────────────────────────────────
    heading1(doc, "2. Estado del Sistema Automatizado")

    heading2(doc, "✅ Funcionando correctamente")
    for item in [
        ("Fix de doble conteo confirmado: ", "0 duplicados por invoice.paid en el período — corrección estable."),
        ("Match por Invoice ID: ", "89 transacciones matcheadas correctamente en el período extendido."),
        ("Match secundario: ", "12 transacciones adicionales por combinación email + monto + timestamp."),
        ("Detección de reembolsos: ", "Eventos charge.refunded capturados correctamente."),
    ]:
        bullet(doc, item[1], bold_prefix=item[0], color=C_GREEN)

    heading2(doc, "⚠️ Puntos de atención activos")
    for item in [
        ("Pagos no capturados post-gap: ",
         "Cobros exitosos en Stripe que no llegaron al sistema por timeouts del endpoint N8N — Stripe agota reintentos."),
        ("Suscripciones canceladas sin datos de cliente: ",
         "Los eventos de cancelación no incluyen email en el payload. Se identifican solo por ID de suscripción."),
    ]:
        bullet(doc, item[1], bold_prefix=item[0], color=C_ORANGE)

    add_divider(doc)

    # ── 3. Cobros Duplicados ─────────────────────────────────────────────────
    heading1(doc, "3. Hallazgo: Cobros Duplicados — Semana del 24–28 de Abril")

    body(doc, (
        "Durante la semana del 24 al 28 de abril se detectaron 5 clientes cobrados dos veces. "
        "Cuando se actualiza manualmente una suscripción en Stripe, el sistema genera un cargo inmediato "
        "de prorrateo adicional al cobro regular del ciclo — si ambos coinciden en el mismo período, "
        "el cliente recibe dos cargos."
    ))

    add_callout(doc,
        "Estos cargos ocurrieron antes del 29 de abril y no aparecen en el reporte manual actual. "
        "Se identificaron a través del reporte automático N8N. Para documentarlos formalmente, "
        "se requiere exportar el reporte Stripe del 24–28 de abril desde el Dashboard.",
        icon="📋"
    )

    add_table(doc,
        ["Cliente", "Monto x cobro", "Tipo de duplicado", "Diferencia"],
        [
            ["moisesmontalvowa@gmail.com", "$85.00 × 2", "Dos actualizaciones de suscripción", "6 minutos"],
            ["cristinamonge@chokolat.com.ec", "$135.00 × 2", "Actualización + ciclo regular", "11.2 horas"],
            ["nmoras12@hotmail.com", "$75.00 × 2", "Actualización + ciclo regular", "18.4 horas"],
            ["sebastianbt@hotmail.es", "$75.00 × 2", "Actualización + ciclo regular", "5.1 horas"],
            ["admin@cmiconstructions.com", "~$85.00 × 2", "Creación duplicada de suscripción", "< 1 minuto"],
        ],
        col_widths=[6.5, 3.5, 5, 3.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "Monto total en cobros duplicados identificados: ", bold=True, color=C_TEXT, size=10)
    add_run(p, "~$855.00", bold=True, color=C_RED, size=10)

    heading2(doc, "Acciones recomendadas")
    for item in [
        "Verificar en Stripe Dashboard si estos cobros fueron reembolsados o están pendientes.",
        "Exportar reporte manual Stripe del 24–28 de abril para confirmar el estado final de cada transacción.",
        "moisesmontalvowa@gmail.com fue cobrado doble y tiene 8 intentos fallidos en una tercera factura — contacto prioritario (ver Sección 5).",
        "Para futuras actualizaciones masivas: aplicar cambios al próximo ciclo en Stripe en lugar de generar cargo inmediato, o hacerlas fuera de ventanas de cobro activas.",
    ]:
        bullet(doc, item)

    add_divider(doc)

    # ── 4. Pagos sin capturar ────────────────────────────────────────────────
    heading1(doc, "4. Pagos Exitosos Sin Capturar en el Sistema")

    body(doc, (
        "Los siguientes pagos aparecen como completados en Stripe pero no fueron registrados por el "
        "sistema automático. Ocurre cuando el endpoint de N8N no responde al webhook y Stripe agota "
        "sus reintentos. Todos requieren backfill manual en el Google Sheet."
    ))

    heading2(doc, "Mayo 1–5 · 5 transacciones — $394.75")
    add_table(doc,
        ["Fecha (UTC)", "Monto", "Cliente", "Invoice ID"],
        [
            ["May 1, 22:05", "$75.00",  "raulsaman@gmail.com",              "in_1TSPBRFNzeVGXAdUJdRBoyjv"],
            ["May 2, 01:55", "$84.75",  "gloriana023@hotmail.com",          "in_1TSRpnFNzeVGXAdUJFAdqdw7"],
            ["May 2, 23:57", "$75.00",  "marianols85@gmail.com",            "in_1TSnPTFNzeVGXAdUDt1WVsZY"],
            ["May 3, 07:52", "$85.00",  "wingzoneswissplaza@gnail.com",     "in_1TStstFNzeVGXAdUjQDirsp6"],
            ["May 4, 21:41", "$75.00",  "contamakimo@gmail.com",            "in_1TTTETFNzeVGXAdU9W4D0u4I"],
        ],
        col_widths=[3, 2.5, 5.5, 7.5],
    )

    heading2(doc, "Abril 29–30 · 9 transacciones — $1,221.05")
    add_table(doc,
        ["Fecha (UTC)", "Monto", "Cliente", "Invoice ID"],
        [
            ["Abr 29, 17:45", "$96.05",   "dondegeorgeguacima@gmail.com",       "in_1TRcAlFNzeVGXAdUCxr3lYkG"],
            ["Abr 29, 18:19", "$75.00",   "somos@fresheatspanama.com",          "in_1TRbkvFNzeVGXAdUnwjStiZZ"],
            ["Abr 29, 20:20", "$85.00",   "facturacionelectronica.chs@agrisal.com", "in_1TRderFNzeVGXAdUSb62TEAj"],
            ["Abr 29, 20:33", "$75.00",   "ajtp223@gmail.com",                  "in_1TRdqxFNzeVGXAdUYrzXFp2B"],
            ["Abr 30, 15:36", "$85.00",   "calderon.fillo@gmail.com",           "in_1TRvgrFNzeVGXAdUh0ZjNfJx"],
            ["Abr 30, 17:52", "$75.00",   "manuel_tandazo16@hotmail.com",       "in_1TRylIFNzeVGXAdUTpH8mQ6o"],
            ["Abr 30, 19:42", "$85.00",   "luis25911@hotmail.com",              "in_1TRzWGFNzeVGXAdUNs9wHIli"],
            ["Abr 30, 21:21", "$570.00 ⚠️","jonathanflores1@gmail.com",          "in_1TS218FNzeVGXAdUZSasBUDT"],
            ["Abr 30, 23:35", "$75.00",   "ejelogistiko.ventas@gmail.com",      "in_1TS395FNzeVGXAdUv2tSjq0k"],
        ],
        col_widths=[3, 2.5, 5.5, 7.5],
    )

    add_callout(doc,
        "El cobro de $570.00 de jonathanflores1@gmail.com del 30 de abril es el mayor monto sin registrar. "
        "Priorizar su verificación en Stripe Webhook Logs.",
        icon="⚠️"
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    add_run(p, "Total pendiente de backfill: ", bold=True, color=C_TEXT, size=10)
    add_run(p, "14 transacciones — $1,615.80", bold=True, color=C_RED, size=10)

    heading2(doc, "Procedimiento de verificación")
    bullet(doc, "Stripe Dashboard → Developers → Webhooks → [endpoint N8N] → Webhook Logs")
    bullet(doc, "Buscar cada Invoice ID para confirmar si el evento fue enviado / si hubo timeout")
    bullet(doc, "Insertar manualmente los registros confirmados en el Google Sheet del reporte automático")

    add_divider(doc)

    # ── 5. Clientes en Riesgo ────────────────────────────────────────────────
    heading1(doc, "5. Clientes en Riesgo de Cancelación")

    body(doc, (
        "El sistema detecta clientes con múltiples intentos de cobro fallidos. Stripe cancela "
        "suscripciones automáticamente al agotar su ventana de reintentos (~9 intentos). "
        "Se recomienda contacto directo para los niveles CRÍTICO y ALTO."
    ))

    heading2(doc, "🔴 CRÍTICO — 9 intentos (cancelación inminente o ya ejecutada)")
    add_table(doc,
        ["Email", "Monto mensual", "Estado"],
        [
            ["drnelsonamador@gmail.com",      "$75.00",  "🔴 CRÍTICO"],
            ["amedina@veronavet.cl",          "$75.00",  "🔴 CRÍTICO"],
            ["gerverehgt020883@icloud.com",   "$85.00",  "🔴 CRÍTICO"],
            ["reneemendez01@gmail.com",       "$75.00",  "🔴 CRÍTICO"],
            ["renechavez11@gmail.com",        "$75.00",  "🔴 CRÍTICO"],
            ["marianols85@gmail.com",         "$75.00",  "🔴 CRÍTICO"],
            ["ipamepc2002@gmail.com",         "$75.00",  "🔴 CRÍTICO"],
            ["robertoe28@hotmail.com",        "$105.00", "🔴 CRÍTICO"],
            ["casco@beautyandthebutcher.com", "$75.00",  "🔴 CRÍTICO"],
        ],
        col_widths=[9, 3.5, 3],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "Valor mensual en riesgo CRÍTICO: ", bold=True, color=C_TEXT, size=10)
    add_run(p, "$715.00", bold=True, color=C_RED, size=10)

    heading2(doc, "🟠 ALTO — 6–8 intentos (acción urgente recomendada)")
    add_table(doc,
        ["Email", "Intentos", "Monto mensual"],
        [
            ["antonioebruno986@gmail.com", "8", "$10.00"],
            ["moisesmontalvowa@gmail.com", "8", "$85.00 ⚠️ cobro doble previo"],
            ["fragallegos@me.com",         "7", "$75.00"],
            ["raquelsteakhouse@gmail.com", "7", "$105.00"],
            ["felipe.rojas7131@gmail.com", "7", "$75.00"],
            ["novasfacturae@gmail.com",    "6", "$75.00"],
            ["danielyancor@icloud.com",    "6", "$75.00"],
        ],
        col_widths=[9, 2.5, 4],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    add_run(p, "Valor mensual en riesgo ALTO: ", bold=True, color=C_TEXT, size=10)
    add_run(p, "$500.00", bold=True, color=C_ORANGE, size=10)

    heading2(doc, "🟡 MEDIO — 4–5 intentos (monitoreo activo)")
    add_table(doc,
        ["Email", "Intentos"],
        [
            ["djarquinm@icloud.com",           "5"],
            ["dondegeorgeguacima@gmail.com",    "5"],
            ["admin@cmiconstructions.com",      "5"],
            ["amendozaw@gmail.com",             "4"],
        ],
        col_widths=[12, 2.5],
    )

    add_callout(doc,
        "moisesmontalvowa@gmail.com: cobrado doble en semana del 24 de abril y actualmente con 8 intentos "
        "fallidos en una tercera factura independiente. Posible relación entre el cobro duplicado y la "
        "resistencia a pagar. Contacto directo prioritario.",
        icon="⚠️"
    )

    heading2(doc, "Razones de fallo más frecuentes")
    add_table(doc,
        ["Razón", "Casos estimados"],
        [
            ["Fondos insuficientes (insufficient_funds)",         "~35"],
            ["Declive genérico (generic_decline)",                "~14"],
            ["Fondos insuficientes — socio (partner_insufficient_funds)", "~7"],
            ["No autorizado (do_not_honor)",                      "~3"],
        ],
        col_widths=[11, 3.5],
    )

    add_divider(doc)

    # ── 6. Acciones Pendientes ───────────────────────────────────────────────
    heading1(doc, "6. Plan de Acciones Pendientes")

    add_table(doc,
        ["Prioridad", "Acción", "Responsable"],
        [
            ["🔴 Inmediata", "Contactar los 9 clientes CRÍTICOS con 9 intentos fallidos", "Devotio Rewards"],
            ["🔴 Inmediata", "Verificar en Stripe Dashboard el estado de los cobros duplicados (Abr 24–28)", "Devotio Rewards"],
            ["🔴 Inmediata", "Exportar reporte Stripe Abr 24–28 para documentar formalmente los dobles cobros", "Devotio Rewards"],
            ["🟠 Esta semana", "Revisar Stripe Webhook Logs para los 14 Invoice IDs no capturados", "Blue Phoenix Lab"],
            ["🟠 Esta semana", "Insertar manualmente los 14 registros en Google Sheet (backfill $1,615.80)", "Blue Phoenix Lab"],
            ["🟠 Esta semana", "Investigar patrón de timeouts en Abr 29–30 (9 pagos sin capturar)", "Blue Phoenix Lab"],
            ["🟡 Próximas semanas", "Implementar enriquecimiento automático de subscription.deleted con datos de cliente", "Blue Phoenix Lab"],
            ["🟡 Próximas semanas", "Definir procedimiento para actualizaciones masivas sin generar cobros duplicados", "Devotio Rewards + Blue Phoenix Lab"],
        ],
        col_widths=[3.5, 9.5, 4.5],
        zebra=True,
    )

    add_divider(doc)

    # ── Footer ───────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(14)
    add_run(footer_p,
        "Reporte generado con script de reconciliación automatizado · Blue Phoenix Lab · Mayo 2026",
        color=C_MUTED, size=8, italic=True
    )

    # Save
    out_path = "output/Devotio_Reporte_Validacion_May1-5_2026.docx"
    doc.save(out_path)
    print(f"✅ Documento guardado: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
